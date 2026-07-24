import sys
import json
from pathlib import Path
import docx
from docx import Document
from docx.table import Table
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# Helpers for cell styling
def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)

def set_cell_margins(cell, top=70, start=90, bottom=70, end=90):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in('w:tcMar')
    if margins is None:
        margins = OxmlElement('w:tcMar')
        tc_pr.append(margins)
    for side, value in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = margins.find(qn('w:' + side))
        if node is None:
            node = OxmlElement('w:' + side)
            margins.append(node)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')

def set_table_borders(table, color='D3D3D3', size='4'):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in('w:tblBorders')
    if borders is None:
        borders = OxmlElement('w:tblBorders')
        tbl_pr.append(borders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        tag = qn('w:' + edge)
        element = borders.find(tag)
        if element is None:
            element = OxmlElement('w:' + edge)
            borders.append(element)
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), size)
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), color)

def no_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        element = OxmlElement('w:' + edge)
        element.set(qn('w:val'), 'nil')
        borders.append(element)
    tbl_pr.append(borders)

def set_table_widths(table, widths_cm):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(width)
                tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
                tc_w = tc_pr.find(qn('w:tcW'))
                if tc_w is None:
                    tc_w = OxmlElement('w:tcW')
                    tc_pr.append(tc_w)
                tc_w.set(qn('w:w'), str(int(width / 2.54 * 1440)))
                tc_w.set(qn('w:type'), 'dxa')

def set_font(run, size=10, bold=False, italic=False, color=None):
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_text(cell, text='', size=10, bold=False, italic=False, color=None, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    paragraph = cell.paragraphs[0]
    paragraph.alignment = alignment
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1
    run = paragraph.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic, color=color)
    return paragraph

def replace_in_paragraphs(paragraphs, data):
    for p in paragraphs:
        text = p.text
        if '{' in text and '}' in text:
            for key, val in data.items():
                placeholder = "{" + str(key) + "}"
                if placeholder in text:
                    replaced = False
                    for run in p.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, str(val))
                            replaced = True
                    if not replaced:
                        p.text = p.text.replace(placeholder, str(val))

def replace_placeholders(doc, data):
    flat_data = {
        'TenKhachHang': data.get('TenKhachHang', ''),
        'MaKH': data.get('MaKH', ''),
        'DiaChi': data.get('DiaChi', ''),
        'SDT': data.get('SDT', ''),
        'NgayLap': data.get('NgayLap', ''),
        'SoPhieu': data.get('SoPhieu', ''),
        'DienGiai': data.get('DienGiai', ''),
        'TongTienHang': data.get('TongTienHang', ''),
        'TienChietKhau': data.get('TienChietKhau', '0'),
        'ChietKhauKhac': data.get('ChietKhauKhac', '0'),
        'TienSauChietKhau': data.get('TienSauChietKhau', ''),
        'TongThanhToan': data.get('TongThanhToan', '')
    }
    
    replace_in_paragraphs(doc.paragraphs, flat_data)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_in_paragraphs(cell.paragraphs, flat_data)

def replace_table_with_matrix(doc, target_table, num_rows, num_cols):
    tbl = target_table._tbl
    parent = tbl.getparent()
    idx = parent.index(tbl)
    
    new_tbl_doc = doc.add_table(rows=num_rows, cols=num_cols)
    new_tbl = new_tbl_doc._tbl
    
    new_tbl.getparent().remove(new_tbl)
    parent.insert(idx, new_tbl)
    parent.remove(tbl)
    
    return Table(new_tbl, doc)


def main():
    if len(sys.argv) < 3:
        print("Usage: python generate_docx_matrix.py <json_path> <output_path>")
        sys.exit(1)

    json_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2])
    
    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    # Resolve template dynamically relative to script location
    template_path = Path(__file__).parent / "samples" / "In Don dat hang.docx"
    if not template_path.exists():
        # Fallback to absolute path if needed
        template_path = Path("d:/Web-Order-Santino/backend-app/samples/In Don dat hang.docx")
        
    if not template_path.exists():
        print(f"Error: Template not found at {template_path}")
        sys.exit(1)
        
    doc = Document(template_path)

    replace_placeholders(doc, data)

    active_sizes = set()
    raw_lines = data.get('ChiTietDonHang', []) or data.get('lines', []) or []
    
    for line in raw_lines:
        sizes = line.get('chi_tiet_size', [])
        if isinstance(sizes, str):
            try:
                sizes = json.loads(sizes)
            except:
                sizes = []
        for s in sizes:
            sz = s.get('size') or s.get('Size') or s.get('ten_size')
            qty = s.get('qty') or s.get('Qty') or s.get('so_luong') or s.get('Quantity')
            try:
                qty = float(qty)
            except:
                qty = 0
            if sz and qty > 0:
                active_sizes.add(str(sz).strip())

    priority = ['38', '39', '40', '41', '42', '43', '44', '45', '46', '47', '48', 
                'S', 'M', 'L', 'XL', '2XL', 'XXL', '3XL', 'XXXL', '4XL', '5XL']
    
    def get_sort_key(s):
        s_str = str(s).strip().upper()
        if s_str in priority:
            return (0, priority.index(s_str))
        try:
            return (1, float(s_str))
        except ValueError:
            return (2, s_str)
            
    sorted_sizes = sorted(list(active_sizes), key=get_sort_key)
    
    target_table = None
    for table in doc.tables:
        found = False
        for row in table.rows:
            for cell in row.cells:
                if "{#ChiTietDonHang}" in cell.text or "ChiTietDonHang" in cell.text:
                    target_table = table
                    found = True
                    break
            if found:
                break
        if found:
            break

    if target_table is None:
        print("Error: Could not find items table containing {#ChiTietDonHang} placeholder.")
        doc.save(output_path)
        sys.exit(0)

    num_sizes = len(sorted_sizes)
    if num_sizes > 0:
        sz_width = 8.0 / num_sizes
    else:
        sz_width = 8.0
        
    widths = [5.5, 1.3] + [sz_width] * num_sizes + [1.3, 2.5]
    num_cols = 4 + num_sizes
    num_rows = 2 + len(raw_lines)
    
    items_tbl = replace_table_with_matrix(doc, target_table, num_rows, num_cols)
    items_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(items_tbl, color='D3D3D3', size='4')
    set_table_widths(items_tbl, widths)
    
    headers = ['SẢN PHẨM', 'MÀU'] + sorted_sizes + ['TỔNG', 'THÀNH TIỀN']
    for idx, text in enumerate(headers):
        cell = items_tbl.rows[0].cells[idx]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cell, 'F2F2F2')
        set_cell_margins(cell, top=80, bottom=80)
        add_text(cell, text, size=8.5, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    size_totals = {sz: 0 for sz in sorted_sizes}
    grand_qty = 0
    
    orange_color = (217, 131, 38)
    gray_color = (120, 120, 120)
    
    for r_idx, line in enumerate(raw_lines):
        row = items_tbl.rows[r_idx + 1]
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell, top=60, bottom=60)
            
        prod_code = line.get('ten_hang_2') or line.get('MaHang') or line.get('ma_hang') or ''
        prod_name = line.get('TenHang') or line.get('ten_hang') or ''
        if ' (Size:' in prod_name:
            prod_name = prod_name.split(' (Size:')[0]
        if prod_code and prod_name.startswith(prod_code):
            prod_name = prod_name[len(prod_code):].strip(' -:')
            
        cell_prod = row.cells[0]
        p = cell_prod.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        
        run_code = p.add_run(prod_code + "\n")
        set_font(run_code, size=9, bold=True, color=orange_color)
        run_name = p.add_run(prod_name)
        set_font(run_name, size=8, color=gray_color)
        
        cell_color = row.cells[1]
        add_text(cell_color, line.get('mau', line.get('MauSac', '')), size=8.5, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        
        sizes = line.get('chi_tiet_size', [])
        if isinstance(sizes, str):
            try:
                sizes = json.loads(sizes)
            except:
                sizes = []
                
        line_sizes = {}
        for s in sizes:
            sz = s.get('size') or s.get('Size') or s.get('ten_size')
            qty = s.get('qty') or s.get('Qty') or s.get('so_luong') or s.get('Quantity')
            try:
                qty = int(float(qty))
            except:
                qty = 0
            if sz:
                line_sizes[str(sz).strip()] = qty
                
        for s_idx, sz in enumerate(sorted_sizes):
            cell_sz = row.cells[2 + s_idx]
            q = line_sizes.get(sz, 0)
            if q > 0:
                add_text(cell_sz, str(q), size=8.5, bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                size_totals[sz] += q
            else:
                add_text(cell_sz, '-', size=8.5, color=gray_color, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                
        qty_total = line.get('so_luong') or line.get('Quantity') or line.get('SoLuong')
        try:
            qty_total = int(float(qty_total))
        except:
            qty_total = sum(line_sizes.values())
            
        grand_qty += qty_total
        cell_tot = row.cells[2 + num_sizes]
        add_text(cell_tot, str(qty_total), size=8.5, bold=True, color=orange_color, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        
        thanh_tien = line.get('ThanhTien') or line.get('thanh_tien') or line.get('Amount') or ''
        try:
            val = float(str(thanh_tien).replace(',', '').replace('.', ''))
            thanh_tien = "{:,.0f}".format(val).replace(',', '.')
        except:
            pass
        cell_money = row.cells[3 + num_sizes]
        add_text(cell_money, str(thanh_tien), size=8.5, bold=True, alignment=WD_ALIGN_PARAGRAPH.RIGHT)

    row_total = items_tbl.rows[-1]
    for cell in row_total.cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_margins(cell, top=60, bottom=60)
        set_cell_shading(cell, 'FAFAFA')
        
    add_text(row_total.cells[0], 'Tổng cộng:', size=8.5, bold=True, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    add_text(row_total.cells[1], '', size=8.5)
    
    for s_idx, sz in enumerate(sorted_sizes):
        cell_sz = row_total.cells[2 + s_idx]
        q = size_totals[sz]
        if q > 0:
            add_text(cell_sz, str(q), size=8.5, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        else:
            add_text(cell_sz, '-', size=8.5, color=gray_color, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            
    add_text(row_total.cells[2 + num_sizes], str(grand_qty), size=8.5, bold=True, color=orange_color, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    
    grand_money = data.get('TongThanhToan', data.get('TongTienHang', ''))
    try:
        val = float(str(grand_money).replace(',', '').replace('.', ''))
        grand_money = "{:,.0f}".format(val).replace(',', '.')
    except:
        pass
    add_text(row_total.cells[3 + num_sizes], str(grand_money), size=8.5, bold=True, color=orange_color, alignment=WD_ALIGN_PARAGRAPH.RIGHT)

    doc.save(output_path)
    print(f"Generated hybrid size matrix DOCX at: {output_path}")


if __name__ == '__main__':
    main()
