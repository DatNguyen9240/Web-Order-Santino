import sys
import os
from pathlib import Path
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUTPUT_SAMPLES = ROOT / 'backend-app' / 'samples' / 'In Don dat hang.docx'
OUTPUT_TEMPLATES = ROOT / 'docx-templates' / 'phieu-dat-hang.docx'
LOGO = ROOT / 'images' / 'logo.png'

DEFAULT_FONT = 'Segoe UI'

def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)

def set_cell_margins(cell, top=30, start=60, bottom=30, end=60):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in('w:tcMar')
    if margins is None:
        margins = OxmlElement('w:tcMar')
        tc_pr.append(margins)
    for side, value in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = margins.find(qn('w:' + side))
        if node is None:
            node = OxmlElement('w:' + side)
            margins.append(node)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')

def set_table_borders(table, color='D0D0D0', size='4'):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in('w:tblBorders')
    if borders is None:
        borders = OxmlElement('w:tblBorders')
        tbl_pr.append(borders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        tag = qn('w:' + edge)
        element = borders.find(tag)
        if element is None:
            element = OxmlElement('w:' + edge)
            borders.append(element)
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), size)
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), color)

def set_table_widths(table, widths_cm):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            row.cells[idx].width = Cm(width)
            tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn('w:tcW'))
            if tc_w is None:
                tc_w = OxmlElement('w:tcW')
                tc_pr.append(tc_w)
            tc_w.set(qn('w:w'), str(int(width / 2.54 * 1440)))
            tc_w.set(qn('w:type'), 'dxa')

def set_font(run, name=DEFAULT_FONT, size=9, bold=False, italic=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:ascii'), name)
    run._element.rPr.rFonts.set(qn('w:hAnsi'), name)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_text(cell_or_para, text='', size=9, bold=False, italic=False, color=None, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    if hasattr(cell_or_para, 'paragraphs'):
        paragraph = cell_or_para.paragraphs[0]
    else:
        paragraph = cell_or_para
    paragraph.alignment = alignment
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.05
    run = paragraph.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic, color=color)
    return paragraph

def no_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        element = OxmlElement('w:' + edge)
        element.set(qn('w:val'), 'nil')
        borders.append(element)
    tbl_pr.append(borders)

def main():
    OUTPUT_SAMPLES.parent.mkdir(parents=True, exist_ok=True)
    OUTPUT_TEMPLATES.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    
    # Page setup: Margins 0.6 cm top/bottom, 0.8 cm left/right for perfect A4 1-page fit
    section = doc.sections[0]
    section.top_margin = Cm(0.6)
    section.bottom_margin = Cm(0.6)
    section.left_margin = Cm(0.8)
    section.right_margin = Cm(0.8)
    section.header_distance = Cm(0.3)
    section.footer_distance = Cm(0.3)

    style = doc.styles['Normal']
    style.font.name = DEFAULT_FONT
    style._element.rPr.rFonts.set(qn('w:ascii'), DEFAULT_FONT)
    style._element.rPr.rFonts.set(qn('w:hAnsi'), DEFAULT_FONT)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), DEFAULT_FONT)
    style.font.size = Pt(9)

    # Printable width: 21.0 - 1.6 = 19.4 cm

    # ==================== 1. HEADER ====================
    header = doc.add_table(rows=1, cols=2)
    header.alignment = WD_TABLE_ALIGNMENT.CENTER
    no_table_borders(header)
    set_table_widths(header, [5.5, 13.9])
    
    logo_cell, company_cell = header.rows[0].cells
    logo_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    company_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(logo_cell, top=0, bottom=0, start=0, end=0)
    set_cell_margins(company_cell, top=0, bottom=0, start=0, end=0)

    if LOGO.exists():
        p_logo = logo_cell.paragraphs[0]
        p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_logo.paragraph_format.space_before = Pt(0)
        p_logo.paragraph_format.space_after = Pt(0)
        p_logo.add_run().add_picture(str(LOGO), width=Cm(3.8))
    else:
        add_text(logo_cell, 'SANTINO', size=20, bold=True, color=(156, 107, 48), alignment=WD_ALIGN_PARAGRAPH.LEFT)

    add_text(company_cell, 'CÔNG TY CP LSP VIỆT NAM', size=11, bold=True, color=(17, 17, 17), alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    
    p = company_cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    set_font(p.add_run('Số 48, Phố Lạc Trung, Q. Hai Bà Trưng, Hà Nội'), size=8, color=(80, 80, 80))
    
    p = company_cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    set_font(p.add_run('Tel: (024) 3204 9988  |  Fax: (024) 3215 1142'), size=8, color=(80, 80, 80))
    
    p = company_cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    set_font(p.add_run('Email: info@santino.com.vn  |  www.santino.com.vn'), size=8, color=(80, 80, 80))

    # ==================== 2. TITLE ====================
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(4)
    p_title.paragraph_format.space_after = Pt(1)
    run_title = p_title.add_run('PHIẾU ĐẶT HÀNG')
    set_font(run_title, size=15, bold=True, color=(17, 17, 17))

    # Sub row: Date left/center, Document No boxed right
    sub_table = doc.add_table(rows=1, cols=2)
    sub_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    no_table_borders(sub_table)
    set_table_widths(sub_table, [12.9, 6.5])
    
    date_cell, doc_no_cell = sub_table.rows[0].cells
    date_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    doc_no_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    add_text(date_cell, '{NgayLap}', size=9, italic=True, color=(70, 70, 70), alignment=WD_ALIGN_PARAGRAPH.CENTER)
    
    # Doc No with nice tan/gold shaded box
    set_cell_shading(doc_no_cell, 'F4EFE4')
    set_cell_margins(doc_no_cell, top=20, bottom=20, start=60, end=60)
    add_text(doc_no_cell, 'Số: {SoPhieu}', size=9, bold=True, color=(17, 17, 17), alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Spacer
    p_sp2 = doc.add_paragraph()
    p_sp2.paragraph_format.space_before = Pt(1)
    p_sp2.paragraph_format.space_after = Pt(1)

    # ==================== 3. CUSTOMER INFO TABLE ====================
    info_table = doc.add_table(rows=3, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(info_table, [12.4, 7.0])
    set_table_borders(info_table, color='E0E0E0', size='4')
    
    # Row 0: KHÁCH HÀNG & MÃ KHÁCH HÀNG
    cell_r0c0 = info_table.rows[0].cells[0]
    cell_r0c1 = info_table.rows[0].cells[1]
    set_cell_margins(cell_r0c0, top=25, bottom=25, start=60, end=60)
    set_cell_margins(cell_r0c1, top=25, bottom=25, start=60, end=60)
    set_cell_shading(cell_r0c0, 'F7F7F5')
    set_cell_shading(cell_r0c1, 'F7F7F5')
    
    p00 = cell_r0c0.paragraphs[0]
    p00.paragraph_format.line_spacing = 1.05
    p00.add_run('KHÁCH HÀNG  ').bold = True
    set_font(p00.runs[0], size=8.5, color=(90, 90, 90))
    p00.add_run('{TenKhachHang}').bold = True
    set_font(p00.runs[1], size=8.5, color=(17, 17, 17))

    p01 = cell_r0c1.paragraphs[0]
    p01.paragraph_format.line_spacing = 1.05
    p01.add_run('MÃ KHÁCH HÀNG  ').bold = True
    set_font(p01.runs[0], size=8.5, color=(90, 90, 90))
    p01.add_run('{MaKH}').bold = True
    set_font(p01.runs[1], size=8.5, color=(17, 17, 17))

    # Row 1: ĐỊA CHỈ & SỐ ĐIỆN THOẠI
    cell_r1c0 = info_table.rows[1].cells[0]
    cell_r1c1 = info_table.rows[1].cells[1]
    set_cell_margins(cell_r1c0, top=25, bottom=25, start=60, end=60)
    set_cell_margins(cell_r1c1, top=25, bottom=25, start=60, end=60)
    set_cell_shading(cell_r1c0, 'F7F7F5')
    set_cell_shading(cell_r1c1, 'F7F7F5')

    p10 = cell_r1c0.paragraphs[0]
    p10.paragraph_format.line_spacing = 1.05
    p10.add_run('ĐỊA CHỈ  ').bold = True
    set_font(p10.runs[0], size=8.5, color=(90, 90, 90))
    p10.add_run('{DiaChi}').bold = False
    set_font(p10.runs[1], size=8.5, color=(17, 17, 17))

    p11 = cell_r1c1.paragraphs[0]
    p11.paragraph_format.line_spacing = 1.05
    p11.add_run('SỐ ĐIỆN THOẠI  ').bold = True
    set_font(p11.runs[0], size=8.5, color=(90, 90, 90))
    p11.add_run('{SDT}').bold = True
    set_font(p11.runs[1], size=8.5, color=(17, 17, 17))

    # Row 2: Merged DIỄN GIẢI across 2 columns
    cell_r2c0 = info_table.rows[2].cells[0]
    cell_r2c1 = info_table.rows[2].cells[1]
    cell_r2_merged = cell_r2c0.merge(cell_r2c1)
    set_cell_margins(cell_r2_merged, top=25, bottom=25, start=60, end=60)
    set_cell_shading(cell_r2_merged, 'FAF6EE')
    
    p2 = cell_r2_merged.paragraphs[0]
    p2.paragraph_format.line_spacing = 1.05
    p2.add_run('DIỄN GIẢI  ').bold = True
    set_font(p2.runs[0], size=8.5, color=(90, 90, 90))
    p2.add_run('{DienGiai}').bold = False
    set_font(p2.runs[1], size=8.5, color=(17, 17, 17))

    # Spacer before main table
    p_sp3 = doc.add_paragraph()
    p_sp3.paragraph_format.space_before = Pt(2)
    p_sp3.paragraph_format.space_after = Pt(2)

    # ==================== 4. MAIN PRODUCTS TABLE ====================
    # 5 Columns: STT (0.9cm), SẢN PHẨM (5.9cm), SIZE × SỐ LƯỢNG (6.8cm), TỔNG (1.8cm), THÀNH TIỀN (4.0cm)
    items_table = doc.add_table(rows=2, cols=5)
    items_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(items_table, [0.9, 5.9, 6.8, 1.8, 4.0])
    set_table_borders(items_table, color='C0C0C0', size='6')

    # Header Row - White Background, Black Bold Text
    headers = ['STT', 'SẢN PHẨM', 'SIZE × SỐ LƯỢNG', 'TỔNG', 'THÀNH TIỀN']
    header_aligns = [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT]
    
    for idx, text in enumerate(headers):
        cell = items_table.rows[0].cells[idx]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cell, 'FFFFFF') # White Background
        set_cell_margins(cell, top=45, bottom=45, start=50, end=50)
        add_text(cell, text, size=8.5, bold=True, color=(17, 17, 17), alignment=header_aligns[idx])

    # Detail Row inside Docxtemplater Loop {#ChiTietDonHang}
    detail_row = items_table.rows[1]
    
    # Cell 0: STT
    cell_stt = detail_row.cells[0]
    cell_stt.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell_stt, top=35, bottom=35, start=20, end=20)
    add_text(cell_stt, '{#ChiTietDonHang}{STT}', size=8.5, bold=True, color=(17, 17, 17), alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Cell 1: SẢN PHẨM (MaHang bold gold/brown, ten_hang_goc, mau)
    cell_prod = detail_row.cells[1]
    cell_prod.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell_prod, top=35, bottom=35, start=60, end=60)
    p_prod = cell_prod.paragraphs[0]
    p_prod.paragraph_format.space_before = Pt(0)
    p_prod.paragraph_format.space_after = Pt(0)
    p_prod.paragraph_format.line_spacing = 1.05
    run_code = p_prod.add_run('{MaHang}\n')
    set_font(run_code, size=9, bold=True, color=(156, 107, 48)) # Brown/Gold Product Code
    run_name = p_prod.add_run('{ten_hang_goc}\n')
    set_font(run_name, size=8, color=(50, 50, 50))
    run_color = p_prod.add_run('Màu: {mau}')
    set_font(run_color, size=8, italic=False, color=(100, 100, 100))

    # Cell 2: SIZE × SỐ LƯỢNG
    cell_size = detail_row.cells[2]
    cell_size.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell_size, top=35, bottom=35, start=60, end=60)
    add_text(cell_size, '{size_qty_text}', size=8.5, color=(17, 17, 17), alignment=WD_ALIGN_PARAGRAPH.LEFT)

    # Cell 3: TỔNG (Quantity)
    cell_qty = detail_row.cells[3]
    cell_qty.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell_qty, top=35, bottom=35, start=20, end=20)
    add_text(cell_qty, '{SoLuong}', size=9, bold=True, color=(17, 17, 17), alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Cell 4: THÀNH TIỀN
    cell_money = detail_row.cells[4]
    cell_money.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell_money, top=35, bottom=35, start=60, end=60)
    add_text(cell_money, '{ThanhTien} đ{/ChiTietDonHang}', size=9, bold=True, color=(17, 17, 17), alignment=WD_ALIGN_PARAGRAPH.RIGHT)

    # ==================== 5. SUMMARY LINE BELOW TABLE ====================
    summary_table = doc.add_table(rows=1, cols=1)
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(summary_table, [19.4])
    set_table_borders(summary_table, color='CCCCCC', size='6')
    
    cell_sum = summary_table.rows[0].cells[0]
    cell_sum.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_shading(cell_sum, 'F4F4F2')
    set_cell_margins(cell_sum, top=35, bottom=35, start=80, end=80)
    
    p_sum = cell_sum.paragraphs[0]
    p_sum.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_sum.paragraph_format.space_before = Pt(0)
    p_sum.paragraph_format.space_after = Pt(0)
    
    r_sum_title = p_sum.add_run('Tổng theo size: ')
    set_font(r_sum_title, size=8.5, bold=True, color=(40, 40, 40))
    
    r_sum_val = p_sum.add_run('{tong_theo_size}')
    set_font(r_sum_val, size=8.5, bold=False, color=(20, 20, 20))

    # Spacer
    p_sp4 = doc.add_paragraph()
    p_sp4.paragraph_format.space_before = Pt(2)
    p_sp4.paragraph_format.space_after = Pt(2)

    # ==================== 6. TOTALS & WORDS SECTION ====================
    totals_table = doc.add_table(rows=4, cols=2)
    totals_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    no_table_borders(totals_table)
    set_table_widths(totals_table, [11.9, 7.5])
    
    for r in totals_table.rows:
        for c in r.cells:
            set_cell_margins(c, top=10, bottom=10)

    # Row 0: Left: Tổng số lượng: 46 sản phẩm | Right: Tổng tiền hàng: 63.270.000 đ
    p_left0 = totals_table.rows[0].cells[0].paragraphs[0]
    p_left0.paragraph_format.line_spacing = 1.05
    r_lbl_qty = p_left0.add_run('Tổng số lượng: ')
    set_font(r_lbl_qty, size=9, bold=True, color=(17, 17, 17))
    r_val_qty = p_left0.add_run('{TongSoLuong} sản phẩm')
    set_font(r_val_qty, size=9, bold=True, color=(17, 17, 17))

    p_right0 = totals_table.rows[0].cells[1].paragraphs[0]
    p_right0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_right0.paragraph_format.line_spacing = 1.05
    r_lbl_th = p_right0.add_run('Tổng tiền hàng: ')
    set_font(r_lbl_th, size=9, bold=True, color=(80, 80, 80))
    r_val_th = p_right0.add_run('{TongTienHang} đ')
    set_font(r_val_th, size=9, bold=True, color=(17, 17, 17))

    # Row 1: Left: Bằng chữ: ... | Right: Chiết khấu: 0 đ
    p_left1 = totals_table.rows[1].cells[0].paragraphs[0]
    p_left1.paragraph_format.line_spacing = 1.05
    r_lbl_bc = p_left1.add_run('Bằng chữ: ')
    set_font(r_lbl_bc, size=8.5, italic=True, color=(80, 80, 80))
    r_val_bc = p_left1.add_run('{TienBangChu}')
    set_font(r_val_bc, size=8.5, italic=True, color=(80, 80, 80))

    p_right1 = totals_table.rows[1].cells[1].paragraphs[0]
    p_right1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_right1.paragraph_format.line_spacing = 1.05
    r_lbl_ck = p_right1.add_run('Chiết khấu: ')
    set_font(r_lbl_ck, size=8.5, color=(80, 80, 80))
    r_val_ck = p_right1.add_run('{TienChietKhau} đ')
    set_font(r_val_ck, size=8.5, bold=True, color=(17, 17, 17))

    # Row 2: Left: (blank) | Right: Chiết khấu khác: 0 đ
    p_right2 = totals_table.rows[2].cells[1].paragraphs[0]
    p_right2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_right2.paragraph_format.line_spacing = 1.05
    r_lbl_ckk = p_right2.add_run('Chiết khấu khác: ')
    set_font(r_lbl_ckk, size=8.5, color=(80, 80, 80))
    r_val_ckk = p_right2.add_run('{ChietKhauKhac} đ')
    set_font(r_val_ckk, size=8.5, bold=True, color=(17, 17, 17))

    # Row 3: Left: (blank) | Right: TỔNG THANH TOÁN: 63.270.000 đ
    p_right3 = totals_table.rows[3].cells[1].paragraphs[0]
    p_right3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_right3.paragraph_format.line_spacing = 1.05
    r_lbl_tt = p_right3.add_run('TỔNG THANH TOÁN:\n')
    set_font(r_lbl_tt, size=9.5, bold=True, color=(17, 17, 17))
    r_val_tt = p_right3.add_run('{TongThanhToan} đ')
    set_font(r_val_tt, size=11.5, bold=True, color=(156, 107, 48))

    # Spacer
    p_sp5 = doc.add_paragraph()
    p_sp5.paragraph_format.space_before = Pt(4)
    p_sp5.paragraph_format.space_after = Pt(2)

    # ==================== 7. SIGNATURES TABLE ====================
    signs = doc.add_table(rows=2, cols=5)
    signs.alignment = WD_TABLE_ALIGNMENT.CENTER
    no_table_borders(signs)
    set_table_widths(signs, [3.88, 3.88, 3.88, 3.88, 3.88])
    
    labels = ['Người nhận', 'Người giao', 'Thủ kho', 'Kế toán', 'Thủ trưởng']
    for idx, label in enumerate(labels):
        cell_top = signs.rows[0].cells[idx]
        cell_top.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_margins(cell_top, top=10, bottom=10)
        add_text(cell_top, label, size=9, bold=True, color=(17, 17, 17), alignment=WD_ALIGN_PARAGRAPH.CENTER)
        
        cell_btm = signs.rows[1].cells[idx]
        cell_btm.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_margins(cell_btm, top=10, bottom=10)
        add_text(cell_btm, '(Ký / họ tên)', size=8, italic=True, color=(120, 120, 120), alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Spacer
    p_sp6 = doc.add_paragraph()
    p_sp6.paragraph_format.space_before = Pt(8)
    p_sp6.paragraph_format.space_after = Pt(2)

    # ==================== 8. FOOTER NOTE ====================
    p_note = doc.add_paragraph()
    p_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_note.paragraph_format.space_before = Pt(2)
    p_note.paragraph_format.space_after = Pt(1)
    r_note = p_note.add_run('Mẫu trình bày gọn "Size × Số lượng" giúp giữ hóa đơn trong một trang A4 mà không thu nhỏ chữ quá mức.')
    set_font(r_note, size=8, italic=True, color=(130, 130, 130))

    # Save to both locations with error handling for locked files
    doc.core_properties.title = 'Phiếu đặt hàng Santino'
    doc.core_properties.subject = 'Mẫu DOCX trình bày gọn theo ma trận size'
    
    saved_paths = []
    for target in [OUTPUT_SAMPLES, OUTPUT_TEMPLATES]:
        try:
            doc.save(target)
            saved_paths.append(str(target))
        except PermissionError:
            print(f'[WARNING] Cannot overwrite: {target}')
            print('   -> File is OPEN in Word. Please CLOSE the file in Word and run again!')
        except Exception as e:
            print(f'[ERROR] Failed saving {target}: {e}')

    if saved_paths:
        print('[SUCCESS] Created DOCX template at:')
        for p in saved_paths:
            print(f' - {p}')

if __name__ == '__main__':
    main()
